from docx import Document

doc = Document("稻鄉.docx")
print(doc.paragraphs)
for para in doc.paragraphs:
    print(para.text)
